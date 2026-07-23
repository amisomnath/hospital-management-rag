"""Load approved TXT, Markdown, PDF and DOCX knowledge documents."""

import hashlib
from dataclasses import dataclass
from pathlib import Path

SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}


@dataclass(slots=True)
class LoadedDocument:
    """Text and metadata extracted from one source file."""

    path: Path
    title: str
    content_type: str
    text: str
    checksum: str


def calculate_checksum(data: bytes) -> str:
    """Return a SHA-256 checksum for change detection."""

    return hashlib.sha256(data).hexdigest()


def load_document(path: Path) -> LoadedDocument:
    """Load TXT, Markdown or PDF content from disk.

    PDF support is imported lazily so the main API can still start when the
    optional PDF dependency has not yet been installed.
    """

    path = path.resolve()
    if not path.exists() or not path.is_file():
        raise ValueError(f"Document does not exist: {path}")

    raw = path.read_bytes()
    suffix = path.suffix.lower()

    if suffix in {".txt", ".md"}:
        text = raw.decode("utf-8", errors="replace")
        content_type = "text/markdown" if suffix == ".md" else "text/plain"
    elif suffix == ".pdf":
        try:
            from pypdf import PdfReader
        except ImportError as exc:  # pragma: no cover - optional dependency
            raise RuntimeError(
                "pypdf is required to ingest PDF files. Install requirements.txt."
            ) from exc

        reader = PdfReader(path)
        pages = []
        for page_number, page in enumerate(reader.pages, start=1):
            page_text = page.extract_text() or ""
            pages.append(f"[Page {page_number}]\n{page_text}")
        text = "\n\n".join(pages)
        content_type = "application/pdf"
    elif suffix == ".docx":
        try:
            from docx import Document
        except ImportError as exc:  # pragma: no cover - dependency guard
            raise RuntimeError(
                "python-docx is required to ingest DOCX files."
            ) from exc
        document = Document(path)
        text = "\n\n".join(
            paragraph.text
            for paragraph in document.paragraphs
            if paragraph.text.strip()
        )
        content_type = (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    else:
        raise ValueError("Only TXT, Markdown, PDF and DOCX files are supported")

    cleaned = "\n".join(line.rstrip() for line in text.splitlines()).strip()
    if not cleaned:
        raise ValueError(f"No readable text was found in {path.name}")

    return LoadedDocument(
        path=path,
        title=path.stem.replace("_", " ").replace("-", " ").title(),
        content_type=content_type,
        text=cleaned,
        checksum=calculate_checksum(raw),
    )
