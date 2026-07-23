"""Build the Bengali beginner handbook for this project."""

from pathlib import Path
from typing import cast

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.styles.style import ParagraphStyle

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "Hospital_RAG_Bangla_Beginner_Guide.docx"


def shade(paragraph, fill: str = "F3F4F6") -> None:
    properties = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def add_code(document: DocumentType, text: str) -> None:
    paragraph = document.add_paragraph()
    shade(paragraph)
    run = paragraph.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)


def add_bullets(document: DocumentType, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_steps(document: DocumentType, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Number")


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.65)
section.bottom_margin = Inches(0.65)
section.left_margin = Inches(0.7)
section.right_margin = Inches(0.7)

styles = doc.styles
normal_style = cast(ParagraphStyle, styles["Normal"])
normal_style.font.name = "Noto Sans Bengali"
normal_style.font.size = Pt(10.5)
for style_name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
    cast(ParagraphStyle, styles[style_name]).font.name = "Noto Sans Bengali"

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Hospital RAG + FastAPI Project\nএকদম সহজ বাংলা Handbook")
run.bold = True
run.font.size = Pt(22)
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.add_run(
    "ক্লাস ৫-এর শিক্ষার্থীকে বোঝানোর মতো ভাষায় • Run • Database • API • "
    "Hugging Face • Groq • প্রতিটি File"
).italic = True

doc.add_heading("১. Project-টি আসলে কী?", level=1)
doc.add_paragraph(
    "ভাবো, এটি একটি ছোট ডিজিটাল হাসপাতাল অফিস এবং একটি প্রশ্ন-উত্তর রোবট। "
    "হাসপাতালের user, patient, doctor, department ও appointment-এর তথ্য database-এ "
    "রাখে। আবার হাসপাতালের নিয়মের text file পড়ে প্রশ্নের কাছাকাছি অংশ খুঁজে উত্তর দেয়।"
)
add_bullets(
    doc,
    [
        "FastAPI = স্কুলের reception desk; কে কোন দরজা দিয়ে আসবে তা ঠিক করে।",
        "SQLAlchemy = reception desk ও database-এর মাঝের দোভাষী।",
        "SQLite/PostgreSQL = তথ্য রাখার আলমারি।",
        "RAG = পুরো বই মুখস্থ না করে আগে দরকারি পাতা খুঁজে নেওয়া।",
        "Embedding = কথাকে সংখ্যার তালিকায় বদলানো, যাতে একই অর্থের কথা মেলানো যায়।",
        "Vector store = সেই সংখ্যাগুলোর আলমারি।",
        "LLM = পাওয়া তথ্য দিয়ে সুন্দর ভাষায় উত্তর বানানো শিক্ষক।",
        "WebSocket = ফোন কলের মতো খোলা connection; বারবার নতুন connection লাগে না।",
    ],
)

doc.add_heading("২. সবচেয়ে সহজে এখনই Run করার নিয়ম", level=1)
doc.add_paragraph("প্রথমবার key ও বড় model ছাড়াই চালানোই সবচেয়ে নিরাপদ। Linux terminal-এ:")
add_code(
    doc,
    "cd /home/user470/Downloads/hospital_rag_fastapi_codebase/hospital_rag_fastapi\n"
    "source .venv/bin/activate\n"
    "unset DEBUG\n"
    "python -m pip install -r requirements.txt",
)
doc.add_paragraph(
    "কেন ‘unset DEBUG’? তোমার computer-এর process environment-এ DEBUG=release থাকলে "
    "সেটি boolean নয় এবং .env-এর DEBUG=true-কে ঢেকে দেয়।"
)
doc.add_paragraph(".env খুলে প্রথম সহজ run-এর জন্য এই মানগুলো রাখো:")
add_code(
    doc,
    "APP_ENV=development\n"
    "DEBUG=true\n"
    "DATABASE_URL=sqlite:///./hospital.db\n"
    "EMBEDDING_BACKEND=hash\n"
    "LLM_PROVIDER=retrieval_only\n"
    "SECRET_KEY=নিজের-একটি-লম্বা-গোপন-লেখা",
)
doc.add_paragraph("তারপর:")
add_code(
    doc,
    "alembic upgrade head\n"
    "python scripts/ingest_knowledge_base.py\n"
    "uvicorn app.main:app --reload",
)
add_bullets(
    doc,
    [
        "Chat page: http://127.0.0.1:8000/",
        "API খেলাঘর (Swagger): http://127.0.0.1:8000/docs",
        "Health check: http://127.0.0.1:8000/api/v1/health",
        "Server বন্ধ: terminal-এ Ctrl+C",
    ],
)

doc.add_heading("৩. Run করলে ভিতরে কী কী ঘটে?", level=1)
add_steps(
    doc,
    [
        "uvicorn app.main:app দেখে app/main.py-এর app object চালু করে।",
        "config.py .env পড়ে settings বানায়।",
        "session.py DATABASE_URL দেখে database connection তৈরি করে।",
        "development mode-এ main.py প্রয়োজনীয় database table তৈরি করে।",
        "router.py সব API দরজা এক জায়গায় জোড়া দেয়।",
        "Browser /docs খুললে FastAPI সব দরজার form দেখায়।",
        "Ingestion script knowledge-base file পড়ে ছোট chunk বানায়।",
        "Embedding service chunk-কে vector বানায়; vector files storage-এ যায়।",
        "প্রশ্ন এলে retriever কাছাকাছি chunk খুঁজে আনে।",
        "Provider retrieval-only/local Hugging Face/Groq নিয়ম অনুযায়ী উত্তর ফেরায়।",
    ],
)

doc.add_heading("৪. Database-এ কী হলো?", level=1)
doc.add_paragraph(
    "SQLite mode-এ project root-এ hospital.db নামে একটি file হয়। Docker mode-এ তথ্য "
    "PostgreSQL container-এর postgres_data volume-এ থাকে।"
)
db_rows = [
    ("users", "login করা মানুষ: email, password hash, role"),
    ("patients", "রোগীর profile ও medical record number"),
    ("doctors", "ডাক্তারের নাম, বিভাগ, licence, availability"),
    ("departments", "Cardiology-এর মতো হাসপাতালের বিভাগ"),
    ("appointments", "কোন patient কখন কোন doctor-এর কাছে যাবে"),
    ("knowledge_documents", "কোন knowledge file নেওয়া হয়েছে"),
    ("knowledge_chunks", "document-এর ছোট ছোট অংশ ও vector id"),
    ("chat_sessions", "একটি কথোপকথনের খাতা"),
    ("chat_messages", "খাতার প্রতিটি user/assistant message"),
]
table = doc.add_table(rows=1, cols=2)
table.style = "Table Grid"
table.rows[0].cells[0].text = "Table"
table.rows[0].cells[1].text = "সহজ কাজ"
for left, right in db_rows:
    cells = table.add_row().cells
    cells[0].text = left
    cells[1].text = right

doc.add_heading("Database নিজের চোখে দেখা", level=2)
add_code(
    doc,
    "sqlite3 hospital.db\n"
    ".tables\n"
    "SELECT id, name FROM departments;\n"
    "SELECT filename, title FROM knowledge_documents;\n"
    ".quit",
)
doc.add_paragraph(
    "sqlite3 command না থাকলে /docs দিয়ে data তৈরি/দেখা যায়, অথবা VS Code-এর SQLite "
    "viewer extension ব্যবহার করা যায়। Password কখনো plain text রাখা হয় না; hashed_password থাকে।"
)

doc.add_heading("৫. RAG-কে গল্প দিয়ে বোঝা", level=1)
doc.add_paragraph(
    "প্রশ্ন: ‘Visiting hour কখন?’ রোবট প্রথমে visiting_hours.txt-এর লেখা ছোট অংশে "
    "ভাগ করে রাখা vector-এর সঙ্গে প্রশ্নের vector মিলায়। সবচেয়ে কাছের ৪টি অংশ নেয়। "
    "তারপর prompt_builder বলে: শুধু পাওয়া তথ্য ধরে উত্তর দাও। এটাই Retrieval-Augmented Generation।"
)
add_bullets(
    doc,
    [
        "data/knowledge_base = অনুমোদিত বইয়ের তাক",
        "chunking.py = বড় পাতা ছোট card-এ কাটে",
        "embedding.py = card-এর অর্থকে number বানায়",
        "vector_store.py = number save/search করে",
        "retriever.py = সবচেয়ে কাছের card আনে",
        "prompt_builder.py = LLM-এর প্রশ্নপত্র বানায়",
        "rag_service.py = পুরো কাজের conductor",
    ],
)

doc.add_heading("৬. Hugging Face-এ কী করতে হবে?", level=1)
doc.add_paragraph(
    "এই project-এর default embedding model sentence-transformers/all-MiniLM-L6-v2 public। "
    "তাই account বা key ছাড়াই প্রথমবার download হতে পারে। প্রথম download-এ internet, সময় ও disk লাগে।"
)
doc.add_heading("Public Hugging Face model চালানো", level=2)
add_code(
    doc,
    "EMBEDDING_BACKEND=sentence_transformers\n"
    "EMBEDDING_MODEL=sentence-transformers/all-MiniLM-L6-v2\n"
    "LLM_PROVIDER=retrieval_only",
)
doc.add_paragraph("তারপর আবার vector বানাও:")
add_code(doc, "python scripts/ingest_knowledge_base.py")
doc.add_heading("Local Hugging Face দিয়ে উত্তর লেখানো", level=2)
add_code(
    doc,
    "LLM_PROVIDER=local_hf\nLOCAL_LLM_MODEL=google/flan-t5-base\nLOCAL_LLM_DEVICE=-1",
)
doc.add_paragraph(
    "-1 মানে CPU। এটি key ছাড়াই চলে, কিন্তু প্রথমবার বড় model download হয় এবং CPU-তে ধীর হতে পারে।"
)
doc.add_heading("কখন Hugging Face token লাগবে?", level=2)
add_steps(
    doc,
    [
        "https://huggingface.co এ account খোলো।",
        "Settings → Access Tokens → New token-এ যাও।",
        "শুধু model পড়তে হলে read বা প্রয়োজনমতো fine-grained token নাও।",
        "Token কাউকে দেখাবে না, screenshot/code/Git-এ দেবে না।",
        "Terminal-এ huggingface-cli login বা hf auth login দিয়ে login করা যায়।",
    ],
)
doc.add_paragraph(
    "Official token guide: https://huggingface.co/docs/hub/security-tokens\n"
    "Model page: https://huggingface.co/sentence-transformers/all-MiniLM-L6-v2"
)

doc.add_heading("৭. Groq key কোথা থেকে এবং কীভাবে?", level=1)
add_steps(
    doc,
    [
        "https://console.groq.com এ account/login করো।",
        "API Keys page (https://console.groq.com/keys) খোলো।",
        "Create API Key চাপো, key copy করো। পরে পুরো key আবার নাও দেখা যেতে পারে।",
        ".env-এ key বসাও; quotation সাধারণত দরকার নেই।",
    ],
)
add_code(
    doc,
    "LLM_PROVIDER=groq\n"
    "GROQ_API_KEY=gsk_তোমার_আসল_key\n"
    "GROQ_MODEL=llama-3.3-70b-versatile",
)
doc.add_paragraph(
    "Official quickstart: https://console.groq.com/docs/quickstart\n"
    "সতর্কতা: .env Git-এ push করবে না। Real patient-এর গোপন তথ্য external LLM-এ পাঠাবে না।"
)

doc.add_heading("৮. API কীভাবে ব্যবহার করবে?", level=1)
doc.add_paragraph(
    "http://127.0.0.1:8000/docs খুলে endpoint বেছে Try it out → Execute করো।"
)
add_bullets(
    doc,
    [
        "POST /auth/register = নতুন user",
        "POST /auth/login = password দেখে access token",
        "GET /users/me = token-এর user",
        "POST/GET /patients = patient তৈরি/দেখা",
        "POST/GET/PATCH /doctors = doctor তৈরি/দেখা/বদলানো",
        "POST/GET /departments = বিভাগ তৈরি/দেখা",
        "POST/GET/PATCH /appointments = appointment তৈরি/দেখা/বদলানো",
        "POST /documents/upload = TXT/MD/PDF upload ও ingest",
        "POST /chat/query = সাধারণ HTTP chat প্রশ্ন",
        "WebSocket /chat/ws/{session_id} = live chat",
        "GET /health = app/database/vector অবস্থার খবর",
    ],
)
doc.add_heading("Register → Login → Token", level=2)
doc.add_paragraph(
    "প্রথমে register। তারপর login response-এর access_token copy করো। Swagger-এর Authorize "
    "button-এ token দিলে protected endpoint বুঝবে তুমি কে। Token হলো অস্থায়ী entry pass।"
)

doc.add_heading("৯. প্রতিটি File-এর কাজ", level=1)
file_notes = {
    "README.md": "মূল পরিচয় ও run command-এর ছোট manual।",
    "requirements.txt": "pip কোন Python package install করবে তার তালিকা।",
    "pyproject.toml": "Poetry metadata এবং Pytest/Ruff/Pyright-এর নিয়ম।",
    "poetry.lock": "Poetry dependency version lock করার file।",
    "setup.py": "Project-কে Python package হিসেবে install করার পুরোনো-style metadata।",
    ".env": "তোমার local secret ও settings; share/commit করা যাবে না।",
    ".env.example": "secret ছাড়া নমুনা settings; copy করে .env বানাতে হয়।",
    ".gitignore": "Git কোন generated/secret file উপেক্ষা করবে।",
    "Dockerfile": "API app-এর Docker image বানানোর recipe।",
    "docker-compose.yml": "API + PostgreSQL একসঙ্গে চালানোর recipe।",
    "alembic.ini": "Alembic migration tool-এর configuration।",
    "alembic/env.py": "settings/database/model metadata Alembic-এর সঙ্গে জোড়া দেয়।",
    "alembic/script.py.mako": "নতুন migration file-এর template।",
    "alembic/versions/7b3be7e8e60b_initial_hospital_schema.py": "প্রথম database schema তৈরি/ফেরত নেওয়ার migration।",
    "app/__init__.py": "app folder-কে Python package বলে চেনায়।",
    "app/main.py": "FastAPI app শুরু, middleware/router/static page জোড়া দেয়।",
    "app/api/deps.py": "প্রতি request-এর DB session ও login করা current user দেয়।",
    "app/api/v1/router.py": "সব version-1 endpoint router একত্র করে।",
    "app/api/v1/endpoints/auth.py": "register ও login দরজা।",
    "app/api/v1/endpoints/users.py": "user list ও নিজের profile দরজা।",
    "app/api/v1/endpoints/patients.py": "patient create/list/read/update দরজা।",
    "app/api/v1/endpoints/doctors.py": "doctor create/list/read/update দরজা।",
    "app/api/v1/endpoints/departments.py": "department create/list দরজা।",
    "app/api/v1/endpoints/appointments.py": "appointment book/list/update দরজা।",
    "app/api/v1/endpoints/documents.py": "knowledge document list/upload/ingest দরজা।",
    "app/api/v1/endpoints/chat.py": "HTTP chat এবং WebSocket chat URL।",
    "app/api/v1/endpoints/health.py": "server, DB ও vector index ঠিক আছে কি না জানায়।",
    "app/core/config.py": ".env পড়ে typed Settings object বানায়।",
    "app/core/logging.py": "terminal log কত বিস্তারিত হবে সেট করে।",
    "app/core/security.py": "password hash/check এবং JWT token create/decode করে।",
    "app/db/base.py": "সব SQLAlchemy model-এর Base এবং model registration।",
    "app/db/session.py": "engine, SessionLocal ও request DB session বানায়।",
    "app/models/user.py": "users table-এর Python নকশা।",
    "app/models/patient.py": "patients table-এর নকশা।",
    "app/models/doctor.py": "doctors table-এর নকশা।",
    "app/models/department.py": "departments table-এর নকশা।",
    "app/models/appointment.py": "appointments table-এর নকশা।",
    "app/models/knowledge_document.py": "knowledge_documents table-এর নকশা।",
    "app/models/knowledge_chunk.py": "knowledge_chunks table-এর নকশা।",
    "app/models/chat_session.py": "chat_sessions table-এর নকশা।",
    "app/models/chat_message.py": "chat_messages table-এর নকশা।",
    "app/models/__init__.py": "সব model export করে, যাতে relation/table জানা থাকে।",
    "app/schemas/auth.py": "register/login input ও token output যাচাইয়ের form।",
    "app/schemas/user.py": "user input/output form।",
    "app/schemas/patient.py": "patient create/update/output form।",
    "app/schemas/doctor.py": "doctor create/update/output form।",
    "app/schemas/department.py": "department create/output form।",
    "app/schemas/appointment.py": "appointment create/update/output form।",
    "app/schemas/document.py": "document ও ingestion response form।",
    "app/schemas/chat.py": "chat message/answer/source/event form।",
    "app/crud/user.py": "user table-এ সরাসরি create/read/list query।",
    "app/crud/patient.py": "patient table-এর database কাজ।",
    "app/crud/doctor.py": "doctor table-এর database কাজ।",
    "app/crud/department.py": "department table-এর database কাজ।",
    "app/crud/appointment.py": "appointment table-এর database কাজ।",
    "app/crud/knowledge_document.py": "document/chunk table-এর database কাজ।",
    "app/crud/chat.py": "chat session/message save করার database কাজ।",
    "app/services/chunking.py": "বড় text overlap রেখে ছোট chunk করে।",
    "app/services/document_loader.py": "TXT/MD/PDF file থেকে text বের করে।",
    "app/services/embedding.py": "text → vector; Hugging Face বা test hash backend।",
    "app/services/vector_store.py": "FAISS/NumPy vector save, load ও nearest search।",
    "app/services/retriever.py": "score অনুযায়ী দরকারি knowledge chunk বেছে নেয়।",
    "app/services/prompt_builder.py": "context ও প্রশ্ন দিয়ে নিরাপদ prompt সাজায়।",
    "app/services/medical_guard.py": "প্রশ্ন medical/hospital scope-এর কি না দেখে।",
    "app/services/emergency_guard.py": "জরুরি বিপদের শব্দ ধরলে emergency guidance দেয়।",
    "app/services/rag_service.py": "ingest, rebuild, retrieve—RAG-এর বড় coordinator।",
    "app/services/chat_service.py": "guard → retrieve → provider → history পুরো chat flow চালায়।",
    "app/llm/base.py": "সব LLM provider-এর একই interface/ফলাফলের নকশা।",
    "app/llm/factory.py": ".env দেখে retrieval/local/Groq provider বেছে নেয়।",
    "app/llm/retrieval_only.py": "LLM ছাড়াই পাওয়া passage উত্তর হিসেবে দেয়।",
    "app/llm/local_hf.py": "নিজের computer-এ Transformers model দিয়ে উত্তর লেখে।",
    "app/llm/groq_provider.py": "Groq API-তে prompt পাঠিয়ে উত্তর আনে।",
    "app/websocket/manager.py": "live WebSocket connection রাখা/বাদ/বার্তা পাঠানো।",
    "app/websocket/chat_handler.py": "WebSocket message validate করে ChatService চালায়।",
    "scripts/ingest_knowledge_base.py": "knowledge folder bulk ingest/rebuild করার CLI।",
    "static/chat.html": "browser chat page-এর কাঠামো।",
    "static/chat.js": "browser থেকে WebSocket connect/send/receive।",
    "static/styles.css": "chat page দেখতে সুন্দর করার style।",
    "data/knowledge_base/admission_policy.txt": "ভর্তি নীতির sample knowledge।",
    "data/knowledge_base/appointment_guide.txt": "appointment-এর sample knowledge।",
    "data/knowledge_base/emergency_information.txt": "emergency sample knowledge।",
    "data/knowledge_base/patient_faq.txt": "patient FAQ sample knowledge।",
    "data/knowledge_base/visiting_hours.txt": "visiting hour sample knowledge।",
    "tests/conftest.py": "isolated test DB/settings ও TestClient fixture।",
    "tests/integration/test_api.py": "অনেক layer একসঙ্গে API ঠিক আছে কি না পরীক্ষা।",
    "tests/integration/test_rag_pipeline.py": "ingest থেকে answer পর্যন্ত পরীক্ষা।",
    "tests/integration/test_websocket.py": "live chat protocol পরীক্ষা।",
    "tests/unit/test_chunking.py": "chunk কাটার ছোট পরীক্ষা।",
    "tests/unit/test_medical_guard.py": "medical scope guard পরীক্ষা।",
    "tests/unit/test_prompt_builder.py": "prompt বানানো পরীক্ষা।",
    "tests/unit/test_retriever.py": "nearest chunk retrieval পরীক্ষা।",
}

all_files = sorted(
    path.relative_to(ROOT).as_posix()
    for path in ROOT.rglob("*")
    if path.is_file()
    and not any(
        part in {".venv", ".ruff_cache", "__pycache__", ".pytest_cache"}
        for part in path.parts
    )
    and path.name not in {"hospital.db", OUTPUT.name}
)
for filename in all_files:
    note = file_notes.get(filename)
    if note is None and filename.endswith("/__init__.py"):
        note = "এই folder-কে Python package হিসেবে চেনায়; সাধারণত glue file।"
    if note is None and filename.startswith("storage/"):
        note = "Ingestion-এর পরে তৈরি হওয়া vector/metadata data file।"
    if note is None and filename.startswith(".vscode/"):
        note = "এই project-এর VS Code editor settings।"
    if note is None:
        note = "Project support/generated file; নাম ও parent folder অনুযায়ী configuration/data রাখে।"
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.add_run(filename + " — ").bold = True
    paragraph.add_run(note)

doc.add_heading("১০. Docker দিয়ে চালানো", level=1)
doc.add_paragraph("Docker installed থাকলে আলাদা PostgreSQL install না করেও:")
add_code(
    doc,
    "docker compose up --build\n"
    "docker compose exec api python scripts/ingest_knowledge_base.py",
)
add_bullets(
    doc,
    [
        "db container = PostgreSQL database",
        "api container = FastAPI app",
        "postgres_data volume = database বাঁচিয়ে রাখে",
        "vector_data volume = vector files বাঁচিয়ে রাখে",
        "hf_cache volume = downloaded Hugging Face model আবার download কমায়",
        "সব মুছতে docker compose down -v দিলে data-ও মুছে যাবে—সাবধানে।",
    ],
)

doc.add_heading("১১. Test এবং code quality", level=1)
add_code(
    doc,
    "source .venv/bin/activate\n"
    "pytest -q\n"
    "ruff check app tests alembic scripts\n"
    "pyright",
)
doc.add_paragraph(
    "বর্তমান যাচাইয়ে ১২টি test pass, Ruff pass এবং Pyright-এ ০ error/০ warning। "
    "Test-এ hash embedding চলে, তাই model download লাগে না।"
)

doc.add_heading("১২. খুব সাধারণ সমস্যা ও সহজ সমাধান", level=1)
problems = [
    ("DEBUG boolean error", "terminal-এ unset DEBUG; .env-এ DEBUG=true/false রাখো।"),
    ("ModuleNotFoundError", "venv activate করে pip install -r requirements.txt চালাও।"),
    ("Port 8000 busy", "পুরোনো server বন্ধ করো বা --port 8001 দাও।"),
    ("Model download ধীর", "প্রথমে EMBEDDING_BACKEND=hash ব্যবহার করো।"),
    ("Groq key missing", "retrieval_only চালাও বা .env-এ valid GROQ_API_KEY দাও।"),
    ("Chat-এ knowledge নেই", "ingest script চালাও এবং data/knowledge_base file দেখো।"),
    ("Database বদল দেখা যায় না", "alembic upgrade head চালাও; প্রয়োজন বুঝে migration বানাও।"),
    (
        "Ruff popup",
        "VS Code Developer: Reload Window; local pyproject config ব্যবহার নিশ্চিত করো।",
    ),
]
problem_table = doc.add_table(rows=1, cols=2)
problem_table.style = "Table Grid"
problem_table.rows[0].cells[0].text = "সমস্যা"
problem_table.rows[0].cells[1].text = "সমাধান"
for problem, solution in problems:
    cells = problem_table.add_row().cells
    cells[0].text = problem
    cells[1].text = solution

doc.add_heading("১৩. শেখার সঠিক ক্রম", level=1)
add_steps(
    doc,
    [
        "প্রথমে hash + retrieval_only দিয়ে run করো।",
        "/docs-এ department, patient, doctor, appointment তৈরি করো।",
        "hospital.db-তে table/data দেখো।",
        "sample knowledge ingest করে chat প্রশ্ন করো।",
        "তারপর sentence_transformers চালিয়ে vector search বোঝো।",
        "তারপর local_hf চেষ্টা করো; computer দুর্বল হলে বাদ দাও।",
        "সবশেষে fake/demo data দিয়ে Groq mode চেষ্টা করো।",
        "একটি file বদলালে test + Ruff + Pyright চালাও।",
    ],
)

doc.add_heading("১৪. কোন File-এর পরে কোন File পড়বে?", level=1)
doc.add_paragraph(
    "Project alphabet অনুযায়ী পড়বে না। একটি request যে রাস্তা দিয়ে যায়, সেই রাস্তা ধরে "
    "এক file থেকে পরের file-এ যাবে। শুরু করবে app/main.py দিয়ে।"
)

doc.add_heading("Application শুরুর রাস্তা", level=2)
add_code(
    doc,
    "uvicorn\n"
    "  ↓\n"
    "app/main.py\n"
    "  ↓\n"
    "app/api/v1/router.py\n"
    "  ↓\n"
    "app/api/v1/endpoints/*.py",
)
add_steps(
    doc,
    [
        "app/main.py পড়ো—FastAPI app, middleware, router ও static page এখানে যুক্ত হয়।",
        "app/api/v1/router.py পড়ো—সব API endpoint এখানে এক জায়গায় জোড়া লাগে।",
        "যে feature বুঝবে, তার endpoint file খোলো। Patient হলে patients.py।",
    ],
)

doc.add_heading("Patient তৈরি হওয়ার রাস্তা", level=2)
add_code(
    doc,
    "endpoints/patients.py\n"
    "  ↓ request নেয়\n"
    "schemas/patient.py\n"
    "  ↓ JSON যাচাই করে\n"
    "crud/patient.py\n"
    "  ↓ database query করে\n"
    "models/patient.py\n"
    "  ↓ table-এর নকশা\n"
    "database",
)
doc.add_paragraph(
    "একই নিয়মে Doctor, Department ও Appointment পড়বে: endpoint → schema → CRUD → model → database।"
)

doc.add_heading("Login ও Token-এর রাস্তা", level=2)
add_code(
    doc,
    "endpoints/auth.py\n"
    "  ↓\n"
    "schemas/auth.py\n"
    "  ↓\n"
    "crud/user.py\n"
    "  ↓\n"
    "core/security.py\n"
    "  ↓\n"
    "models/user.py → database",
)
add_bullets(
    doc,
    [
        "auth.py register/login request নেয়।",
        "schemas/auth.py email ও password form যাচাই করে।",
        "crud/user.py user তৈরি বা খোঁজে।",
        "security.py password hash/check এবং JWT token তৈরি/decode করে।",
        "api/deps.py token থেকে current user বের করে protected endpoint-কে দেয়।",
    ],
)

doc.add_heading("Database connection-এর রাস্তা", level=2)
add_code(
    doc,
    ".env-এর DATABASE_URL\n"
    "  ↓\n"
    "core/config.py\n"
    "  ↓\n"
    "db/session.py\n"
    "  ↓\n"
    "api/deps.py\n"
    "  ↓\n"
    "endpoint-এর db parameter\n"
    "  ↓\n"
    "CRUD → Model → Database table",
)

doc.add_heading("Knowledge ingest-এর রাস্তা", level=2)
add_code(
    doc,
    "data/knowledge_base/*.txt\n"
    "  ↓\n"
    "scripts/ingest_knowledge_base.py\n"
    "  ↓\n"
    "services/rag_service.py\n"
    "  ↓\n"
    "document_loader.py → chunking.py → embedding.py → vector_store.py",
)

doc.add_heading("Chat প্রশ্নের রাস্তা", level=2)
add_code(
    doc,
    "Browser / API question\n"
    "  ↓\n"
    "endpoints/chat.py\n"
    "  ↓\n"
    "websocket/chat_handler.py (WebSocket হলে)\n"
    "  ↓\n"
    "services/chat_service.py\n"
    "  ├→ emergency_guard.py\n"
    "  ├→ medical_guard.py\n"
    "  ├→ rag_service.py → retriever.py → vector_store.py\n"
    "  ├→ prompt_builder.py\n"
    "  └→ llm/factory.py\n"
    "        ├→ retrieval_only.py\n"
    "        ├→ local_hf.py\n"
    "        └→ groq_provider.py",
)
doc.add_paragraph(
    "কোন provider চলবে তা .env-এর LLM_PROVIDER বলে। কোন embedding চলবে তা "
    "EMBEDDING_BACKEND বলে। তাই কোনো flow বোঝার আগে .env দেখবে।"
)

doc.add_heading("File কার সঙ্গে connected—নিজে কীভাবে জানবে?", level=2)
add_steps(
    doc,
    [
        "File-এর উপরের import দেখো। যেমন ‘from app.services.rag_service import RAGService’ মানে এই file rag_service ব্যবহার করছে।",
        "কোনো class/function-এর উপর Ctrl চেপে click বা F12 চাপো—তার definition file খুলবে।",
        "Shift+F12 বা right-click → Find All References দিলে কে কে এটি ব্যবহার করছে দেখাবে।",
        "Ctrl+Shift+F দিয়ে পুরো project-এ RAGService, create_patient বা get_settings-এর মতো নাম search করো।",
        'Terminal-এ rg "RAGService" app অথবা rg "create_patient" app চালানো যায়।',
    ],
)

doc.add_heading("মনে রাখার দুটি Formula", level=2)
add_code(
    doc,
    "সাধারণ API:\n"
    "Request → Endpoint → Schema → Service/CRUD → Model → Database → Response\n\n"
    "Chat/RAG:\n"
    "Question → Safety check → Vector search → Knowledge → Prompt → Provider → Answer",
)

doc.add_heading("১৫. নিরাপত্তার শেষ কথা", level=1)
doc.add_paragraph(
    "এটি শেখার project, চিকিৎসা দেওয়ার machine নয়। জরুরি অবস্থায় স্থানীয় emergency service/doctor-এর "
    "সাহায্য নিতে হবে। Real patient name, phone, report, diagnosis বা অন্য ব্যক্তিগত তথ্য Groq/Hugging "
    "Face-এর মতো external service-এ অনুমতি ও যথাযথ নিরাপত্তা ছাড়া পাঠানো যাবে না।"
)

doc.add_paragraph()
end = doc.add_paragraph()
end.alignment = WD_ALIGN_PARAGRAPH.CENTER
end.add_run("শেষ — আগে সবচেয়ে সহজ run করো, তারপর এক একটি নতুন জিনিস যোগ করো।").bold = True

doc.save(str(OUTPUT))
print(OUTPUT)
